import os
import uuid
import logging
import json
from datetime import datetime, timedelta
from functools import wraps
import tempfile
import hashlib
import re
from pathlib import Path

from flask import Flask, render_template, request, jsonify, session, redirect, url_for
from flask_socketio import SocketIO, emit, disconnect
from flask_cors import CORS
import requests
from dotenv import load_dotenv
from werkzeug.utils import secure_filename
import pandas as pd
import chardet
from PyPDF2 import PdfReader
from docx import Document

# Load environment variables
load_dotenv()

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Flask app configuration
app = Flask(__name__, static_folder='static', static_url_path='')
app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', os.urandom(24).hex())
app.config['SESSION_COOKIE_SECURE'] = os.environ.get('FLASK_ENV') == 'production'
app.config['SESSION_COOKIE_HTTPONLY'] = True
app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'
app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(hours=24)

# File upload configuration
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size
ALLOWED_EXTENSIONS = {'xlsx', 'xls', 'csv', 'pdf', 'docx', 'doc', 'txt'}
UPLOAD_FOLDER = tempfile.gettempdir()

# CORS configuration
CORS(app, resources={r"/api/*": {"origins": "*"}})

# Patch DNS resolution for eventlet
import eventlet
eventlet.monkey_patch()

# SocketIO configuration
socketio = SocketIO(
    app,
    cors_allowed_origins="*",
    async_mode='eventlet',
    ping_timeout=120,
    ping_interval=25,
    max_http_buffer_size=1000000
)

# Admin credentials from environment
ADMIN_USERNAME = os.environ.get('ADMIN_USERNAME', 'admin')
ADMIN_PASSWORD = os.environ.get('ADMIN_PASSWORD', 'secure_password_123')

# Claude API URL - Use host.docker.internal when running in Docker
CLAUDE_API_URL = os.environ.get('CLAUDE_API_URL', 'http://host.docker.internal:8000')

# Active sessions tracking
active_sessions = {}


def allowed_file(filename):
    """Check if file extension is allowed"""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def sanitize_filename(filename):
    """Sanitize filename to prevent security issues"""
    # Get the file extension
    name, ext = os.path.splitext(filename)
    # Remove any path components
    name = os.path.basename(name)
    # Remove special characters
    name = re.sub(r'[^\w\s-]', '', name)
    name = re.sub(r'[-\s]+', '-', name)
    # Add timestamp hash for uniqueness
    timestamp = datetime.utcnow().strftime('%Y%m%d%H%M%S')
    hash_suffix = hashlib.md5(f"{name}{timestamp}".encode()).hexdigest()[:8]
    return f"{name}_{hash_suffix}{ext}"


def detect_encoding(file_path):
    """Detect file encoding"""
    with open(file_path, 'rb') as f:
        raw_data = f.read()
        result = chardet.detect(raw_data)
        return result['encoding'] or 'utf-8'


def convert_excel_to_text(file_path):
    """Convert Excel file to text format"""
    try:
        # Read Excel file
        df = pd.read_excel(file_path, sheet_name=None)
        
        text_output = []
        
        # Process each sheet
        for sheet_name, sheet_df in df.items():
            text_output.append(f"=== Sheet: {sheet_name} ===\n")
            
            # Convert to string representation
            text_output.append(sheet_df.to_string(index=False))
            text_output.append("\n\n")
        
        return '\n'.join(text_output)
    except Exception as e:
        logger.error(f"Error converting Excel file: {e}")
        return f"Error converting Excel file: {str(e)}"


def convert_csv_to_text(file_path):
    """Convert CSV file to text format"""
    try:
        # Detect encoding
        encoding = detect_encoding(file_path)
        
        # Read CSV file
        df = pd.read_csv(file_path, encoding=encoding)
        
        # Convert to string representation
        return df.to_string(index=False)
    except Exception as e:
        logger.error(f"Error converting CSV file: {e}")
        return f"Error converting CSV file: {str(e)}"


def convert_pdf_to_text(file_path):
    """Convert PDF file to text format"""
    try:
        reader = PdfReader(file_path)
        text_output = []
        
        for i, page in enumerate(reader.pages):
            text_output.append(f"=== Page {i + 1} ===\n")
            text_output.append(page.extract_text())
            text_output.append("\n\n")
        
        return '\n'.join(text_output)
    except Exception as e:
        logger.error(f"Error converting PDF file: {e}")
        return f"Error converting PDF file: {str(e)}"


def convert_docx_to_text(file_path):
    """Convert DOCX file to text format"""
    try:
        doc = Document(file_path)
        text_output = []
        
        for paragraph in doc.paragraphs:
            text_output.append(paragraph.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = '\t'.join(cell.text for cell in row.cells)
                text_output.append(row_text)
        
        return '\n'.join(text_output)
    except Exception as e:
        logger.error(f"Error converting DOCX file: {e}")
        return f"Error converting DOCX file: {str(e)}"


def convert_file_to_text(file_path, filename):
    """Convert various file formats to text"""
    ext = filename.rsplit('.', 1)[1].lower()
    
    if ext in ['xlsx', 'xls']:
        return convert_excel_to_text(file_path)
    elif ext == 'csv':
        return convert_csv_to_text(file_path)
    elif ext == 'pdf':
        return convert_pdf_to_text(file_path)
    elif ext == 'docx':
        return convert_docx_to_text(file_path)
    elif ext == 'doc':
        # For older .doc files, we'll return a message since python-docx doesn't support them
        return "注意: .doc形式（古いWord形式）は直接変換できません。.docx形式で保存し直してからアップロードしてください。"
    elif ext == 'txt':
        encoding = detect_encoding(file_path)
        with open(file_path, 'r', encoding=encoding) as f:
            return f.read()
    else:
        return "Unsupported file format"


def login_required(f):
    """Decorator to check if user is logged in"""
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user_id' not in session:
            if request.path.startswith('/api/'):
                return jsonify({'error': 'Authentication required'}), 401
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated_function


@app.route('/')
def index():
    """Main page - redirect to login or chat"""
    if 'user_id' in session:
        return redirect(url_for('chat'))
    return redirect(url_for('login'))


@app.route('/login')
def login():
    """Login page"""
    return render_template('login.html')


@app.route('/chat')
@login_required
def chat():
    """Chat interface page"""
    return render_template('chat.html', username=session.get('username', 'Admin'))


@app.route('/health')
def health():
    """Health check endpoint"""
    return jsonify({'status': 'healthy', 'timestamp': datetime.utcnow().isoformat()})


@app.route('/api/login', methods=['POST'])
def api_login():
    """Handle login API request"""
    data = request.get_json()
    username = data.get('username', '')
    password = data.get('password', '')
    
    if username == ADMIN_USERNAME and password == ADMIN_PASSWORD:
        session['user_id'] = str(uuid.uuid4())
        session['username'] = username
        session.permanent = True
        
        logger.info(f"User {username} logged in successfully")
        return jsonify({'success': True, 'username': username})
    
    logger.warning(f"Failed login attempt for username: {username}")
    return jsonify({'success': False, 'error': 'Invalid credentials'}), 401


@app.route('/api/logout', methods=['POST'])
@login_required
def api_logout():
    """Handle logout API request"""
    username = session.get('username', 'Unknown')
    session.clear()
    logger.info(f"User {username} logged out")
    return jsonify({'success': True})


@app.route('/api/upload', methods=['POST'])
@login_required
def api_upload():
    """Handle file upload and convert to text"""
    try:
        # Check if file is in request
        if 'file' not in request.files:
            return jsonify({'success': False, 'error': 'No file provided'}), 400
        
        file = request.files['file']
        
        # Check if file is selected
        if file.filename == '':
            return jsonify({'success': False, 'error': 'No file selected'}), 400
        
        # Check if file is allowed
        if not allowed_file(file.filename):
            return jsonify({
                'success': False, 
                'error': f'Invalid file type. Allowed types: {", ".join(ALLOWED_EXTENSIONS)}'
            }), 400
        
        # Sanitize filename
        original_filename = file.filename
        safe_filename = sanitize_filename(file.filename)
        
        # Save file temporarily
        temp_path = os.path.join(UPLOAD_FOLDER, safe_filename)
        file.save(temp_path)
        
        logger.info(f"File uploaded: {original_filename} -> {safe_filename}")
        
        try:
            # Convert file to text
            text_content = convert_file_to_text(temp_path, safe_filename)
            
            # Clean up temporary file
            os.unlink(temp_path)
            
            # Log success
            logger.info(f"File converted successfully: {original_filename}")
            
            return jsonify({
                'success': True,
                'filename': original_filename,
                'text': text_content,
                'length': len(text_content)
            })
            
        except Exception as e:
            # Clean up temporary file on error
            if os.path.exists(temp_path):
                os.unlink(temp_path)
            raise e
            
    except Exception as e:
        logger.error(f"File upload error: {e}")
        return jsonify({
            'success': False,
            'error': f'File processing error: {str(e)}'
        }), 500


@app.route('/api/health', methods=['GET'])
def health_check():
    """Health check endpoint"""
    try:
        # Check Claude API availability
        claude_status = "unavailable"
        try:
            response = requests.get(f"{CLAUDE_API_URL}/status", timeout=5)
            if response.status_code == 200:
                claude_status = "available"
        except:
            pass
        
        return jsonify({
            'status': 'healthy',
            'timestamp': datetime.utcnow().isoformat(),
            'claude_api': claude_status,
            'active_sessions': len(active_sessions)
        })
    except Exception as e:
        logger.error(f"Health check failed: {e}")
        return jsonify({'status': 'unhealthy', 'error': str(e)}), 500


@socketio.on('connect')
def handle_connect():
    """Handle WebSocket connection"""
    logger.info(f"WebSocket connection attempt from {request.remote_addr}")
    logger.info(f"Session data: {dict(session)}")
    logger.info(f"Request SID: {request.sid}")
    
    # Check authentication manually for WebSocket
    if 'user_id' not in session:
        logger.warning(f"WebSocket connection rejected - no session: {request.sid}")
        return False
    
    session_id = session.get('user_id')
    if session_id:
        active_sessions[request.sid] = {
            'user_id': session_id,
            'username': session.get('username', 'Unknown'),
            'connected_at': datetime.utcnow()
        }
        logger.info(f"WebSocket connected successfully: {request.sid} for user {session.get('username')}")
        emit('connected', {'status': 'Connected to WebAI'})
    else:
        logger.error("Session ID exists but is empty")
        disconnect()


@socketio.on('disconnect')
def handle_disconnect():
    """Handle WebSocket disconnection"""
    if request.sid in active_sessions:
        del active_sessions[request.sid]
        logger.info(f"WebSocket disconnected: {request.sid}")


@socketio.on('ping')
def handle_ping(data):
    """Handle ping for testing"""
    logger.info(f"Received ping: {data}")
    emit('pong', {'timestamp': data.get('timestamp'), 'server_time': datetime.utcnow().isoformat()})

@socketio.on('message')
def handle_message(data):
    """Handle incoming chat message"""
    logger.info(f"[MESSAGE] Received message event from {request.sid}")
    logger.info(f"[MESSAGE] Data received: {data}")
    logger.info(f"[MESSAGE] Session info: user_id={session.get('user_id')}, username={session.get('username')}")
    
    # Check authentication
    if 'user_id' not in session:
        logger.error(f"[MESSAGE] No user_id in session for {request.sid}")
        emit('error', {'error': 'Authentication required'})
        return
    
    try:
        user_message = data.get('message', '')
        model = data.get('model', 'claude-3-sonnet')
        web_search = data.get('web_search', True)
        session_id = session.get('user_id')
        
        if not user_message:
            logger.error("[MESSAGE] Empty message received")
            emit('error', {'error': 'Empty message'})
            return
        
        logger.info(f"[MESSAGE] Processing: message='{user_message[:50]}...', model={model}, web_search={web_search}")
        
        # Emit acknowledgment
        emit('message_received', {'status': 'Processing your request...'})
        
        # Prepare request for Claude API
        claude_request = {
            'message': user_message,
            'model': model,
            'stream': True,
            'web_search': web_search
        }
        
        # Send request to Claude API
        try:
            # Check API availability first
            logger.info(f"Checking Claude API health at {CLAUDE_API_URL}/health")
            health_check = requests.get(f"{CLAUDE_API_URL}/health", timeout=5)
            if health_check.status_code != 200:
                logger.error(f"Claude API health check failed: {health_check.status_code}")
                emit('error', {'error': 'Claude API is not available'})
                return
            logger.info("Claude API health check passed")
            
            # Send message to Claude API with model and web_search parameters
            api_request = {
                'content': user_message,
                'model': model,
                'web_search': web_search
            }
            
            logger.info(f"[CLAUDE] Sending request to Claude API: {api_request}")
            logger.info(f"[CLAUDE] URL: {CLAUDE_API_URL}/message")
            
            response = requests.post(
                f"{CLAUDE_API_URL}/message",
                json=api_request,
                stream=True,
                timeout=120,
                headers={'Accept': 'application/x-ndjson'}
            )
            
            logger.info(f"Claude API response status: {response.status_code}")
            if response.status_code != 200:
                logger.error(f"Claude API error response: {response.text}")
                emit('error', {'error': f'Claude API error: {response.status_code}'})
                return
            
            # Stream response chunks
            current_content = ""
            chunk_count = 0
            logger.info("[STREAM] Starting to stream response chunks")
            
            for line in response.iter_lines():
                if line:
                    try:
                        line_str = line.decode('utf-8')
                        logger.debug(f"[STREAM] Raw line {chunk_count}: {line_str[:100]}")
                        
                        # Parse JSON lines
                        if line_str.strip():
                            data = json.loads(line_str)
                            logger.debug(f"[STREAM] Parsed data: {data}")
                            
                            if 'content' in data:
                                chunk_content = data['content']
                                if current_content:
                                    current_content += "\n" + chunk_content
                                else:
                                    current_content = chunk_content
                                
                                chunk_count += 1
                                logger.info(f"[STREAM] Emitting chunk {chunk_count}: {len(chunk_content)} chars to {request.sid}")
                                emit('stream_chunk', {'chunk': current_content})
                                
                            elif 'error' in data:
                                logger.error(f"[STREAM] Error from Claude API: {data['error']}")
                                emit('error', {'error': data['error']})
                                return
                                
                            elif 'status' in data and data['status'] == 'complete':
                                logger.info(f"[STREAM] Stream complete after {chunk_count} chunks")
                                emit('stream_complete', {'status': 'Response complete'})
                                break
                                
                    except json.JSONDecodeError as e:
                        logger.error(f"[STREAM] JSON decode error: {e}, line: {line_str}")
                    except Exception as e:
                        logger.error(f"[STREAM] Error parsing chunk: {e}")
            
            logger.info(f"[STREAM] Final: Sent {chunk_count} chunks, total content length: {len(current_content)}")
            emit('stream_complete', {'status': 'Response complete'})
            
        except requests.exceptions.Timeout:
            logger.error("[CLAUDE] Request timed out")
            emit('error', {'error': 'Request timed out'})
        except requests.exceptions.ConnectionError as e:
            logger.error(f"[CLAUDE] Connection error: {e}")
            emit('error', {'error': 'Claude API is not available'})
        except Exception as e:
            logger.error(f"[CLAUDE] Request failed: {type(e).__name__}: {e}")
            emit('error', {'error': f'Failed to connect to Claude API: {str(e)}'})
            
    except Exception as e:
        logger.error(f"[MESSAGE] Error handling message: {type(e).__name__}: {e}")
        import traceback
        logger.error(f"[MESSAGE] Traceback: {traceback.format_exc()}")
        emit('error', {'error': 'Internal server error'})


@app.errorhandler(404)
def not_found(error):
    """Handle 404 errors"""
    return jsonify({'error': 'Not found'}), 404


@app.errorhandler(500)
def internal_error(error):
    """Handle 500 errors"""
    logger.error(f"Internal server error: {error}")
    return jsonify({'error': 'Internal server error'}), 500


if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    debug = os.environ.get('FLASK_ENV') == 'development'
    
    logger.info(f"Starting WebAI on port {port} (debug={debug})")
    socketio.run(app, host='0.0.0.0', port=port, debug=debug)